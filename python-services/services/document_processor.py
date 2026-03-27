"""
Document processing service for text extraction and chunking
"""

import os
import logging
from typing import List, Dict, Any
import PyPDF2
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
import tiktoken

from models import DocumentChunk

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document text extraction and chunking"""
    
    def __init__(self, max_chunk_size: int = 1000, chunk_overlap: int = 200):
        self.max_chunk_size = max_chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=max_chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=self._count_tokens,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Initialize tokenizer for token counting
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken"""
        return len(self.tokenizer.encode(text))
    
    async def extract_text(self, file_path: str) -> str:
        """Extract text from various document formats"""
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return await self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return await self._extract_from_docx(file_path)
            elif file_extension in ['.txt', '.md']:
                return await self._extract_from_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1} from {file_path}: {e}")
                        continue
                        
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            raise
        
        if not text.strip():
            raise ValueError("No text could be extracted from the PDF")
        
        return text.strip()
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from the DOCX file")
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise
    
    async def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text.strip():
                raise ValueError("The text file is empty")
            
            return text.strip()
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return text.strip()
            except Exception as e:
                logger.error(f"Error reading text file {file_path} with latin-1 encoding: {e}")
                raise
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            raise
    
    async def chunk_text(self, text: str, title: str, source: str) -> List[DocumentChunk]:
        """Split text into chunks for embedding"""
        
        if not text.strip():
            raise ValueError("Cannot chunk empty text")
        
        try:
            # Clean and preprocess text
            cleaned_text = self._clean_text(text)
            
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(cleaned_text)
            
            # Create DocumentChunk objects
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                if chunk_text.strip():  # Only include non-empty chunks
                    chunk = DocumentChunk(
                        content=chunk_text.strip(),
                        chunk_index=i,
                        title=title,
                        source=source,
                        metadata={
                            "token_count": self._count_tokens(chunk_text),
                            "character_count": len(chunk_text),
                            "chunk_size": len(chunk_text.split())
                        }
                    )
                    chunks.append(chunk)
            
            if not chunks:
                raise ValueError("No valid chunks could be created from the text")
            
            logger.info(f"Created {len(chunks)} chunks from text (total tokens: {self._count_tokens(text)})")
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        
        # Remove empty lines and combine
        cleaned_lines = []
        for line in lines:
            if line:
                cleaned_lines.append(line)
        
        # Join with single newlines
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive spaces
        import re
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    def get_document_stats(self, text: str) -> Dict[str, Any]:
        """Get statistics about a document"""
        
        return {
            "character_count": len(text),
            "word_count": len(text.split()),
            "token_count": self._count_tokens(text),
            "line_count": len(text.split('\n')),
            "estimated_chunks": max(1, self._count_tokens(text) // self.max_chunk_size)
        }
